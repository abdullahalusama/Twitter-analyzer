
import streamlit as st
import tweepy
import snscrape.modules.twitter as sntwitter
import pandas as pd
from textblob import TextBlob
import matplotlib.pyplot as plt
import io
import base64
import time
from docx import Document

# Twitter API credentials
API_KEY = "KA3VcqIQfRr1EbcU4PTddw4EM"
API_SECRET = "mOVcaaOeUhYMHu7ow80dROSGwmIE5TPbVhMgKIIIG8PWGcOc8C"
BEARER_TOKEN = "AAAAAAAAAAAAAAAAAAAAALIUxgEAAAAALYV06nmzg7WprA9Wp55ekWpusM8%3DHP6L2IQnkqnbIVhSwsmkm8RvaRyFlsyV4XgRKMEXPtUq1uJQIy"

# Authenticate with Tweepy
client = tweepy.Client(bearer_token=BEARER_TOKEN)

# Function to fetch tweets using Twitter API
def fetch_tweets_api(topic, max_tweets):
    tweets_list = []
    query = f"{topic} -is:retweet lang:en"
    try:
        for tweet in tweepy.Paginator(client.search_recent_tweets, query=query, max_results=10, tweet_fields=["author_id"]).flatten(limit=max_tweets):
            tweets_list.append({
                "username": tweet.author_id,
                "tweet": tweet.text
            })
            time.sleep(1)  # Add delay to avoid hitting rate limits
    except tweepy.TooManyRequests:
        return None  # Return None if rate limit is reached
    return tweets_list

# Function to fetch tweets using snscrape
def fetch_tweets_scrape(topic, max_tweets):
    tweets_list = []
    query = f'{topic} lang:en'
    for i, tweet in enumerate(sntwitter.TwitterSearchScraper(query).get_items()):
        if i >= max_tweets:
            break
        tweets_list.append({
            "username": tweet.username,
            "tweet": tweet.content
        })
    return tweets_list

# Sentiment analysis function
def analyze_sentiment(tweets):
    sentiments = {"Positive": 0, "Neutral": 0, "Negative": 0}
    sentiment_details = []

    for tweet_data in tweets:
        polarity = TextBlob(tweet_data["tweet"]).sentiment.polarity
        if polarity > 0:
            sentiment = "Positive"
            sentiments["Positive"] += 1
        elif polarity < 0:
            sentiment = "Negative"
            sentiments["Negative"] += 1
        else:
            sentiment = "Neutral"
            sentiments["Neutral"] += 1
        sentiment_details.append({
            "username": tweet_data["username"],
            "tweet": tweet_data["tweet"],
            "sentiment": sentiment
        })

    return sentiments, pd.DataFrame(sentiment_details)

# Function to generate pie chart
def generate_pie_chart(sentiments):
    labels = sentiments.keys()
    sizes = sentiments.values()
    colors = ['#66b3ff', '#99ff99', '#ff9999']
    explode = (0.1, 0, 0)

    plt.figure(figsize=(7, 7))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors, explode=explode)
    plt.title('Sentiment Analysis Results')

    # Save plot to a BytesIO object
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    return img

# Function to generate a Word document
def generate_word_file(data):
    doc = Document()
    doc.add_heading("Tweet Sentiment Analysis", level=1)
    for item in data:
        doc.add_paragraph(f"Username: {item['username']}")
        doc.add_paragraph(f"Tweet: {item['tweet']}")
        doc.add_paragraph(f"Sentiment: {item['sentiment']}")
        doc.add_paragraph("\n")  # Blank line for separation

    # Save the Word file to a BytesIO object
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app layout
st.title("Tweet Sentiment Analysis")
st.markdown("Analyze tweets based on a topic and get sentiment distribution!")

# Input form
with st.form("tweet_form"):
    topic = st.text_input("Enter a topic to search for tweets (e.g., AI, Movies):", "")
    max_tweets = st.slider("Number of tweets to fetch:", 10, 100, 50)
    submitted = st.form_submit_button("Analyze")

if submitted:
    if topic.strip() == "":
        st.error("Please enter a valid topic!")
    else:
        st.write(f"Fetching tweets for **{topic}**...")

        # Fetch tweets
        tweets = fetch_tweets_api(topic, max_tweets)
        if tweets is None or len(tweets) == 0:  # If API fails or no tweets fetched
            tweets = fetch_tweets_scrape(topic, max_tweets)

        if not tweets:
            st.error("No tweets found. Please try another topic.")
        else:
            # Perform sentiment analysis
            sentiments, sentiment_df = analyze_sentiment(tweets)

            # Generate and display pie chart
            chart = generate_pie_chart(sentiments)
            st.image(chart, caption="Sentiment Analysis Results", use_container_width=True)

            # Display sentiment details
            st.subheader("Detailed Sentiment Analysis")
            st.dataframe(sentiment_df)

            # Downloadable CSV
            csv = sentiment_df.to_csv(index=False)
            b64_csv = base64.b64encode(csv.encode()).decode()
            href_csv = f'<a href="data:file/csv;base64,{b64_csv}" download="sentiment_analysis.csv">Download CSV</a>'
            st.markdown(href_csv, unsafe_allow_html=True)

            # Downloadable Word file
            word_file = generate_word_file(sentiment_df.to_dict(orient="records"))
            b64_word = base64.b64encode(word_file.getvalue()).decode()
            href_word = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_word}" download="sentiment_analysis.docx">Download Word File</a>'
            st.markdown(href_word, unsafe_allow_html=True)

